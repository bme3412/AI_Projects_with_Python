from flask import Flask, request, send_file, render_template
from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, Document
from llama_index.core.node_parser import SentenceSplitter, SemanticSplitterNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, AIMessage
import os
import pandas as pd
import re
from docx import Document as DocxDocument
from docx.shared import Inches
from io import BytesIO
import PyPDF2
import zipfile

app = Flask(__name__)
os.environ["OPENAI_API_KEY"] = ""

embed_model = OpenAIEmbedding()
splitter = SemanticSplitterNodeParser(
    buffer_size=1,
    breakpoint_percentile_threshold=90,
    embed_model=embed_model
)

# also baseline splitter
base_splitter = SentenceSplitter(chunk_size=512)


def process_text_file(file):
    file_content = file.read().decode('utf-8')
    return [Document(file_content)]


def process_word_file(file):
    document = DocxDocument(file)
    file_content = "\n".join(
        [paragraph.text for paragraph in document.paragraphs])
    return [Document(text=file_content)]


def process_pdf_file(file):
    pdf_reader = PyPDF2.PdfReader(file)
    file_content = "\n".join([page.extract_text()
                             for page in pdf_reader.pages])
    return [Document(text=file_content)]


def extract_info_from_title(title):
    # Extract company name (assuming it's the first part of the title before '_')
    company_name_match = re.search(r'^(.*?)_', title)
    company_name = company_name_match.group(1) if company_name_match else None

    # Extract quarter
    quarter_match = re.search(r'(Q\d)', title)
    quarter = quarter_match.group(1) if quarter_match else None

    # Extract year
    year_match = re.search(r'(\d{4})', title)
    year = year_match.group(1) if year_match else None

    return company_name, quarter, year


@app.route('/', methods=['GET', 'POST'])
def process_file():
    if request.method == 'POST':
        # Get the uploaded file
        file = request.files['file']

        # Determine the file type based on the file extension
        file_extension = os.path.splitext(file.filename)[1].lower()

        # Process the file based on its type
        if file_extension == '.txt':
            documents = process_text_file(file)
        elif file_extension == '.docx':
            documents = process_word_file(file)
        elif file_extension == '.pdf':
            documents = process_pdf_file(file)
        else:
            return "Invalid file type. Only .txt, .docx, and .pdf files are allowed."

        # Extract company name, quarter, and year from the title
        title = documents[0].text.split("\n")[0]
        company_name, quarter, year = extract_info_from_title(title)

        # Split documents into nodes
        nodes = splitter.get_nodes_from_documents(documents)

        # Create an empty DataFrame to store the nodes, summaries, and categories
        df = pd.DataFrame(columns=["Raw_Text", "Summary", "Category"])

        # Initialize the LLM
        llm = ChatOpenAI(model_name="gpt-4-turbo")

        # Generate summary and category for each node in sequential order
        for i, node in enumerate(nodes, start=1):
            summary_query = f"Could you summarize the following text? Return your response which covers the key points and does not miss anything important, please. No need to start with 'The text discusses', etc.\n\n{node.text}"
            summary_result = llm([HumanMessage(content=summary_query)])
            summary = summary_result.content

            category_query = f"Please provide a short category or topic for the following summary:\n\n{summary}"
            category_result = llm([HumanMessage(content=category_query)])
            category = category_result.content

            # Create a new DataFrame for the current node
            node_df = pd.DataFrame({"Raw_Text": [node.text], "Summary": [
                                   summary], "Category": [category]})

            # Append the node DataFrame to the main DataFrame
            df = pd.concat([df, node_df], ignore_index=True)

      

        # Create a BytesIO object to store the CSV file
        csv_buffer = BytesIO()
        df.to_csv(csv_buffer, index=False)
        csv_buffer.seek(0)

        # Create a new Word document
        document = DocxDocument()

        # Add a table to the document
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Write the column headers
        header_row = table.rows[0]
        for i, column_name in enumerate(df.columns):
            header_row.cells[i].text = column_name

        # Write the data rows
        for _, row in df.iterrows():
            new_row = table.add_row()
            for i, value in enumerate(row):
                new_row.cells[i].text = str(value)

        # Create a BytesIO object to store the Word document
        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)

        # Create a zip file containing the CSV and Word files
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            zip_file.writestr(
                f"{company_name}_{quarter}_{year}.csv", csv_buffer.getvalue())
            zip_file.writestr(
                f"{company_name}_{quarter}_{year}.docx", docx_buffer.getvalue())
        zip_buffer.seek(0)

   
        # Create the 'static' directory if it doesn't exist
        os.makedirs('static', exist_ok=True)

        # Render the template with the processed data
        return render_template('result.html', df=df, company_name=company_name, quarter=quarter, year=year)

    return render_template('index.html')


if __name__ == '__main__':
    app.run()
